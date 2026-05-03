import json
import sys
from docx import Document


def main():
    json_path = sys.argv[1]
    docx_path = sys.argv[2]

    with open(json_path, "r", encoding="utf-8") as handle:
        project = json.load(handle)

    document = Document()
    document.add_heading(project["title"], 0)
    document.add_paragraph(project.get("promise", ""))
    if project.get("tableOfContents"):
        document.add_heading("Table des matières", level=1)
        document.add_paragraph(project["tableOfContents"])

    for chapter in project.get("chapters", []):
        document.add_heading(chapter["title"], level=1)
        if chapter.get("summary"):
            document.add_paragraph(chapter["summary"])
        if chapter.get("content"):
            document.add_paragraph(chapter["content"])

    document.save(docx_path)


if __name__ == "__main__":
    main()
